from __future__ import annotations

import io
import math
import re
import unicodedata
from typing import Any

import fitz
import pytesseract
from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

RTL_RE = re.compile(r"[\u0590-\u08FF]")


def norm(s: str) -> str:
    s = unicodedata.normalize("NFC", s or "")
    s = s.replace("ي", "ی").replace("ى", "ی").replace("ك", "ک")
    s = s.replace("\u200e", "").replace("\u200f", "")
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


def is_rtl(s: str) -> bool:
    return bool(RTL_RE.search(s or ""))


def set_run_font(run, font_name: str | None):
    if not font_name:
        return
    run.font.name = font_name
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def set_paragraph_direction(p, rtl: bool):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    ppr = p._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if rtl and bidi is None:
        ppr.append(OxmlElement("w:bidi"))
    elif not rtl and bidi is not None:
        ppr.remove(bidi)


def set_cell_shading(cell, fill="E9EEF7"):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def quality(words: list[dict[str, Any]]) -> float:
    if not words:
        return 0.0
    chars = "".join(w["text"] for w in words)
    useful = sum(bool(re.search(r"[\w\u0590-\u08FF]", w["text"])) for w in words)
    printable = sum(c.isprintable() for c in chars)
    return min(1.0, 0.55 * printable / max(1, len(chars)) + 0.45 * useful / len(words))


def extract_pdf_words(page) -> list[dict[str, Any]]:
    words: list[dict[str, Any]] = []
    data = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)
    for block in data.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                text = norm(span.get("text", ""))
                if not text:
                    continue
                x0, y0, x1, y1 = span["bbox"]
                words.append({
                    "text": text, "x": float(x0), "y": float(y0),
                    "w": max(1.0, float(x1 - x0)), "h": max(1.0, float(y1 - y0)),
                    "font": span.get("font", ""), "size": float(span.get("size", 11)),
                    "flags": int(span.get("flags", 0)), "confidence": 100.0,
                    "block": tuple(block.get("bbox", (0, 0, 0, 0))),
                })
    return words


def ocr_page(page, language: str, quality_mode: str = "balanced"):
    scale = {"high": 3.0, "balanced": 2.4, "fast": 1.7}.get(quality_mode, 2.4)
    psm = "--oem 1 --psm 3" if quality_mode != "fast" else "--oem 1 --psm 6"
    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
    try:
        data = pytesseract.image_to_data(img, lang=language, output_type=pytesseract.Output.DICT, config=psm)
    except pytesseract.TesseractError as exc:
        raise RuntimeError("OCR فارسی/انگلیسی اجرا نشد؛ Tesseract و بسته‌های fas/eng را بررسی کنید.") from exc
    words: list[dict[str, Any]] = []
    confidences: list[float] = []
    for i, raw in enumerate(data.get("text", [])):
        text = norm(raw)
        try:
            conf = float(data.get("conf", [0])[i])
        except Exception:
            conf = 0.0
        if not text or conf < 5:
            continue
        confidences.append(conf)
        words.append({
            "text": text,
            "x": float(data["left"][i]) / scale,
            "y": float(data["top"][i]) / scale,
            "w": max(1.0, float(data["width"][i]) / scale),
            "h": max(1.0, float(data["height"][i]) / scale),
            "font": "", "size": max(8.0, float(data["height"][i]) / scale), "flags": 0,
            "confidence": conf,
            "block": (0, 0, 0, 0),
        })
    return words, (sum(confidences) / len(confidences) if confidences else 0.0)


def group_lines(words: list[dict[str, Any]]) -> list[dict[str, Any]]:
    groups: list[dict[str, Any]] = []
    for w in sorted(words, key=lambda z: (z["y"], z["x"])):
        best = None
        best_delta = 9999.0
        for g in groups[-8:]:
            tolerance = max(3.5, min(12.0, max(g["h"], w["h"]) * 0.55))
            delta = abs(g["y"] - w["y"])
            if delta <= tolerance and delta < best_delta:
                best, best_delta = g, delta
        if best is None:
            best = {"y": w["y"], "h": w["h"], "x0": w["x"], "x1": w["x"] + w["w"], "words": []}
            groups.append(best)
        best["words"].append(w)
        best["h"] = max(best["h"], w["h"])
        best["x0"] = min(best["x0"], w["x"])
        best["x1"] = max(best["x1"], w["x"] + w["w"])
    out = []
    for g in sorted(groups, key=lambda z: z["y"]):
        rtl_score = sum(1 for w in g["words"] if is_rtl(w["text"]))
        rtl = rtl_score >= max(1, math.ceil(len(g["words"]) / 2))
        g["words"].sort(key=lambda x: x["x"], reverse=rtl)
        out.append(g | {"rtl": rtl, "text": " ".join(w["text"] for w in g["words"])})
    return out


def detect_table(lines: list[dict[str, Any]], mode: str = "balanced"):
    if mode == "off" or len(lines) < 3:
        return None
    if mode == "strict" and len(lines) < 4:
        return None
    anchors: dict[int, int] = {}
    for line in lines:
        seen = set()
        for w in line["words"]:
            x = round(w["x"] / 16) * 16
            if x not in seen:
                anchors[x] = anchors.get(x, 0) + 1
                seen.add(x)
    threshold = max(3, math.ceil(len(lines) * (0.55 if mode == "strict" else 0.45)))
    cols = sorted(x for x, n in anchors.items() if n >= threshold)
    if len(cols) < 2 or len(cols) > 12:
        return None
    rows = []
    nonempty = 0
    for line in lines:
        row = []
        for i, x in enumerate(cols):
            left = -1e9 if i == 0 else (cols[i - 1] + x) / 2
            right = 1e9 if i == len(cols) - 1 else (x + cols[i + 1]) / 2
            value = " ".join(w["text"] for w in line["words"] if left <= w["x"] < right).strip()
            row.append(value)
        if sum(bool(v) for v in row) >= 2:
            nonempty += 1
            rows.append(row)
    minimum = 3 if mode == "strict" else 2
    if nonempty < minimum:
        return None
    density = nonempty / max(1, len(lines))
    if density < 0.6:
        return None
    return rows


def add_line(doc: Document, line: dict[str, Any], previous_y: float | None, page_width: float, smart: bool, preserve_style: bool):
    p = doc.add_paragraph()
    set_paragraph_direction(p, line["rtl"])
    if smart and previous_y is not None:
        gap = max(0.0, line["y"] - previous_y)
        p.paragraph_format.space_before = Pt(min(18, gap * 0.28))
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    for idx, w in enumerate(line["words"]):
        text = w["text"] + (" " if idx < len(line["words"]) - 1 else "")
        run = p.add_run(text)
        if preserve_style:
            run.font.size = Pt(max(6, min(72, w.get("size", 11))))
            set_run_font(run, w.get("font"))
            flags = int(w.get("flags", 0))
            run.bold = bool(flags & 16)
            run.italic = bool(flags & 2)
        else:
            run.font.size = Pt(11)
    # approximate original horizontal placement as a Word indent.
    if smart and page_width > 0:
        indent = max(0.0, min(4.0, line["x0"] / 72.0))
        if line["rtl"]:
            p.paragraph_format.right_indent = Inches(indent)
        else:
            p.paragraph_format.left_indent = Inches(indent)
    return p


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_paragraph_direction(p, is_rtl(value))
            run = p.add_run(value)
            run.font.size = Pt(10.5)
            if r == 0:
                run.bold = True
                set_cell_shading(cell)
    return table


def page_size_inches(page):
    # PDF points -> inches, with safe Word section bounds.
    return max(5.0, min(22.0, page.rect.width / 72.0)), max(5.0, min(22.0, page.rect.height / 72.0))


def convert_document(data: bytes, filename: str, ocr_mode="auto", language="fas+eng", table_mode="balanced", ocr_quality="balanced", preserve_style=True, smart_paragraphs=True, safe_tables=True):
    if data[:4] == b"%PDF":
        pdf = fitz.open(stream=data, filetype="pdf")
    else:
        try:
            image = Image.open(io.BytesIO(data)).convert("RGB")
        except Exception as exc:
            raise ValueError("فایل تصویر معتبر نیست.") from exc
        pdf = fitz.open()
        page = pdf.new_page(width=image.width, height=image.height)
        page.insert_image(page.rect, stream=data)

    doc = Document()
    stats = {"pages": pdf.page_count, "ocr_pages": 0, "words": 0, "tables": 0, "confidence": 0.0}
    confidences: list[float] = []

    for pi, page in enumerate(pdf):
        if pi == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()
        width_in, height_in = page_size_inches(page)
        section.page_width = Inches(width_in)
        section.page_height = Inches(height_in)
        section.top_margin = Inches(min(0.65, height_in * 0.08))
        section.bottom_margin = Inches(min(0.65, height_in * 0.08))
        section.left_margin = Inches(min(0.65, width_in * 0.08))
        section.right_margin = Inches(min(0.65, width_in * 0.08))

        words = extract_pdf_words(page)
        native_quality = quality(words)
        use_ocr = ocr_mode == "always" or (ocr_mode == "auto" and (len(words) < 5 or native_quality < 0.72))
        if use_ocr:
            words, conf = ocr_page(page, language, ocr_quality)
            stats["ocr_pages"] += 1
            confidences.append(conf)
        else:
            conf = 100.0
            confidences.append(conf)

        stats["words"] += len(words)
        lines = group_lines(words)
        table = detect_table(lines, table_mode) if safe_tables else None
        if table:
            add_table(doc, table)
            stats["tables"] += 1
        else:
            previous_y = None
            for line in lines:
                add_line(doc, line, previous_y, page.rect.width, smart_paragraphs, preserve_style)
                previous_y = line["y"] + line["h"]

        if pi < pdf.page_count - 1:
            doc.add_page_break()

    stats["confidence"] = sum(confidences) / len(confidences) if confidences else 0.0
    if not doc.paragraphs and not doc.tables:
        p = doc.add_paragraph("متنی از فایل قابل استخراج نبود.")
        set_paragraph_direction(p, True)
    out = io.BytesIO()
    doc.save(out)
    pdf.close()
    return out.getvalue(), stats
